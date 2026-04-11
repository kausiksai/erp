"""Generate a shareable Word document (and optional PDF) with full data
quality findings from Phase 3 validation.

Produces three files under email_automation/logs/:
    phase3_data_quality_report.docx    — shareable Word doc
    phase3_data_quality_report.pdf     — only if docx2pdf / Word is available
    phase3_issues_data.json            — raw finding data for any future use

Run with:
    python -m email_automation.phase3_generate_doc

Takes ~10 minutes for 2,033 invoices (the validation engine is re-run in
read-only mode to guarantee freshness).
"""

from __future__ import annotations

import json
import sys
import time
from collections import defaultdict
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from .db import close_pool, get_conn, get_cursor
from .logger import setup_logging
from .phase3_issues_report import ERROR_DESCRIPTIONS
from .validation.engine import run_full_validation


REPORT_DIR = Path(__file__).resolve().parent / "logs"
REPORT_DIR.mkdir(parents=True, exist_ok=True)
DOCX_PATH = REPORT_DIR / "phase3_data_quality_report.docx"
JSON_PATH = REPORT_DIR / "phase3_issues_data.json"
PDF_PATH = REPORT_DIR / "phase3_data_quality_report.pdf"


# ---------------------------------------------------------------------------
# Priority order (drives the outline of the doc)
# ---------------------------------------------------------------------------
CATEGORY_ORDER = [
    ("Missing reference data",           "Reference data gaps",         "🔴"),
    ("Data quality",                     "Source data errors",          "🔴"),
    ("GST data quality",                 "GST compliance risks",        "🟠"),
    ("Price drift",                      "Price mismatch (PO vs invoice)", "🟠"),
    ("Shortfall / over-billing",         "Over-billing (quantity)",     "🟡"),
    ("Shortfall / over-billing (amount)","Over-billing (amount)",       "🟡"),
    ("Shortfall / partial delivery",     "Partial delivery",            "🟡"),
    ("Shortfall (receipt coverage)",     "GRN coverage gap",            "🟡"),
    ("Cumulative over-billing",          "Cumulative qty over PO",      "🟡"),
    ("Cumulative over-billing (amount)", "Cumulative amount over PO",   "🟡"),
    ("Exception approval",               "PO already fulfilled",        "🟢"),
    ("Open PO data gap",                 "Open PO: missing reference",  "🟡"),
    ("Open PO data quality",             "Open PO: data mismatch",      "🟡"),
]

BUCKET_DESCRIPTIONS = {
    "validated":                 "Clean pass — ready for payment flow",
    "waiting_for_re_validation": "Shortfall — routed to debit note approval",
    "exception_approval":        "PO already fulfilled — routed to exception approval",
    "waiting_for_validation":    "Data errors or missing reference — retries on each run",
}


def collect_data() -> Dict[str, Any]:
    """Run the validation engine in read-only mode over all email_automation
    invoices and return the collected finding data."""
    with get_cursor(readonly=True) as cur:
        cur.execute(
            "SELECT invoice_id FROM invoices WHERE source='email_automation' ORDER BY invoice_id"
        )
        ids = [r["invoice_id"] for r in cur.fetchall()]
    total = len(ids)
    print(f"Running validation on {total} invoices (read-only)...")

    t0 = time.time()
    per_code: Dict[str, List[Dict[str, Any]]] = defaultdict(list)
    bucket_counts: Dict[str, int] = defaultdict(int)
    invoice_error_codes: Dict[int, List[str]] = defaultdict(list)

    for idx, invoice_id in enumerate(ids):
        try:
            with get_conn(readonly=True) as conn:
                result = run_full_validation(conn, invoice_id)
            for err in result.errors:
                per_code[err.code].append(
                    {
                        "invoice_id": invoice_id,
                        "po_id": result.po_id,
                        "message": err.message,
                        "data": {k: str(v) for k, v in (err.data or {}).items()},
                        "line_seq": err.line_seq,
                    }
                )
                invoice_error_codes[invoice_id].append(err.code)
            bucket_counts[result.target_status] += 1
        except Exception as exc:
            print(f"  [warn] invoice {invoice_id} raised {exc}")
        if (idx + 1) % 250 == 0:
            print(f"  progress: {idx + 1}/{total} ({time.time() - t0:.0f}s)")

    print(f"done in {time.time() - t0:.0f}s")

    # Enrich all affected invoices with header info (bulk fetch)
    all_ids = sorted(invoice_error_codes.keys())
    enriched: Dict[int, Dict[str, Any]] = {}
    if all_ids:
        with get_cursor(readonly=True) as cur:
            cur.execute(
                """
                SELECT i.invoice_id, i.invoice_number, i.unit, i.po_number,
                       i.total_amount, i.tax_amount, i.invoice_date,
                       s.supplier_name, s.state_code
                FROM invoices i
                LEFT JOIN suppliers s ON s.supplier_id = i.supplier_id
                WHERE i.invoice_id = ANY(%s)
                """,
                (all_ids,),
            )
            for r in cur.fetchall():
                row = dict(r)
                # serialize non-JSON types
                for k, v in list(row.items()):
                    if hasattr(v, "isoformat"):
                        row[k] = v.isoformat()
                    elif v is not None and not isinstance(v, (str, int, float, bool)):
                        row[k] = str(v)
                enriched[r["invoice_id"]] = row

    return {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "total_invoices": total,
        "bucket_counts": dict(bucket_counts),
        "per_code": {k: v for k, v in per_code.items()},
        "enriched": enriched,
    }


# ---------------------------------------------------------------------------
# DOCX builder
# ---------------------------------------------------------------------------
def _set_cell_text(cell, text: str, bold: bool = False, size: int = 9, color: str = None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _add_shaded_row(table, cells: List[str]):
    row = table.add_row()
    for i, text in enumerate(cells):
        _set_cell_text(row.cells[i], text, bold=True, size=9, color="FFFFFF")
        tcPr = row.cells[i]._tc.get_or_add_tcPr()
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F4E79")
        tcPr.append(shd)


def _format_amt(v):
    if v in (None, ""):
        return "?"
    try:
        return f"₹{float(v):,.2f}"
    except (TypeError, ValueError):
        return str(v)


def build_docx(data: Dict[str, Any]) -> Path:
    doc = Document()

    # -- Page margins ---------------------------------------------------------
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # -- Cover block ----------------------------------------------------------
    title = doc.add_heading("Data Quality Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Bill Register — Automated Validation Findings")
    run.italic = True
    run.font.size = Pt(13)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Source: Bill Register Mar-26.xlsx  |  "
        f"Generated: {data['generated_at']}  |  "
        f"Total invoices: {data['total_invoices']:,}"
    ).font.size = Pt(10)

    doc.add_paragraph()

    # -- 1. Executive summary -------------------------------------------------
    doc.add_heading("1. Executive Summary", level=1)
    total = data["total_invoices"]
    buckets = data["bucket_counts"]

    p = doc.add_paragraph()
    p.add_run(
        f"A total of {total:,} invoices were ingested from the Bill Register "
        f"export and run through the automated validation engine. The engine "
        f"applies 30+ checks covering reference-data integrity, PO matching, "
        f"price enforcement, GST compliance, shortfall detection, and cumulative "
        f"over-billing.\n\n"
    )
    p.add_run("Validation outcome by bucket:").bold = True

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Light Grid Accent 1"
    _add_shaded_row(tbl, ["Bucket", "Count", "%", "Meaning"])
    for bucket, count in sorted(buckets.items(), key=lambda x: -x[1]):
        pct = (count / total * 100) if total else 0
        row = tbl.add_row()
        _set_cell_text(row.cells[0], bucket)
        _set_cell_text(row.cells[1], f"{count:,}")
        _set_cell_text(row.cells[2], f"{pct:.1f}%")
        _set_cell_text(row.cells[3], BUCKET_DESCRIPTIONS.get(bucket, ""))

    doc.add_paragraph()

    # -- 2. Top priorities ----------------------------------------------------
    doc.add_heading("2. Top 4 Priorities for the srimukha Data Team", level=1)

    p = doc.add_paragraph()
    p.add_run(
        "Fixing the four items below will unblock the majority of failing "
        "invoices. Each item is ranked by impact."
    )

    per_code = data["per_code"]
    enriched = data["enriched"]

    def _count(code: str) -> int:
        return len({e["invoice_id"] for e in per_code.get(code, [])})

    priorities = [
        (
            "Priority 1 — Subcontract orders missing from PO export",
            _count("E003_PO_NOT_FOUND"),
            "Invoices reference PO numbers (prefix 'SC') that are not present in PO.xls.",
            "Add SC-prefixed subcontract orders to the daily PO export OR send them as a separate daily email (e.g. 'Notification - Subcontract Order').",
        ),
        (
            "Priority 2 — Invoices posted without any PO reference",
            _count("E002_NO_PO_LINK"),
            "The 'PO / SCO No.' column in the Bill Register is empty for these invoices, so the automated engine cannot link them to any PO.",
            "Make 'PO / SCO No.' a mandatory field when posting invoices in the srimukha ERP. Every invoice row must carry a PO or subcontract-order reference.",
        ),
        (
            "Priority 3 — GST type (intra-state vs inter-state) applied incorrectly",
            _count("E034_INTRA_STATE_WITH_IGST") + _count("E035_INTER_STATE_WITH_CGST_SGST"),
            "Invoices where supplier and place-of-supply are in the same state but IGST is charged, OR in different states but CGST/SGST is charged. This is a compliance risk under GST law.",
            "Supplier / finance — correct the tax head on these invoices. The supplier GSTIN's first two digits must match place_of_supply for CGST+SGST; any mismatch must use IGST.",
        ),
        (
            "Priority 4 — Invoice rate or amount does not match PO",
            _count("E022_LINE_RATE_MISMATCH") + _count("E023_LINE_PRICE_MISMATCH"),
            "Line-level rate drift. Invoice rate differs from PO unit_cost × (1 - disc%), or the line assessable value does not match qty × effective PO rate.",
            "Finance review per case: (a) rate renegotiated but PO not amended, (b) supplier billing above agreed rate, or (c) PO data error (we observed BHARAT PETROLEUM POs with unit_cost stored as ₹1). Issue debit notes or amend POs as appropriate.",
        ),
    ]

    for title, count, what, action in priorities:
        doc.add_heading(title, level=2)
        p = doc.add_paragraph()
        p.add_run("Affected invoices: ").bold = True
        p.add_run(f"{count:,}")
        p = doc.add_paragraph()
        p.add_run("What is happening: ").bold = True
        p.add_run(what)
        p = doc.add_paragraph()
        p.add_run("Action required: ").bold = True
        p.add_run(action)

    # -- 3. Detailed findings by category ------------------------------------
    doc.add_page_break()
    doc.add_heading("3. Detailed Findings by Category", level=1)

    p = doc.add_paragraph()
    p.add_run(
        "Each category below lists all error codes, their plain-English "
        "description, severity, required action, and five real examples with "
        "bill number, supplier, PO, and amount so the data team can investigate "
        "the exact invoices. "
    )
    run = p.add_run(
        "Finding counts are line-level; one invoice can produce multiple "
        "findings. The 'affected invoices' column shows the distinct invoice "
        "count."
    )
    run.italic = True

    # Build a category -> codes index
    code_to_cat: Dict[str, str] = {}
    for code in per_code.keys():
        code_to_cat[code] = ERROR_DESCRIPTIONS.get(code, {}).get("category", "Other")
    by_cat: Dict[str, List[str]] = defaultdict(list)
    for code, cat in code_to_cat.items():
        by_cat[cat].append(code)

    for cat, heading, color in CATEGORY_ORDER:
        if cat not in by_cat:
            continue
        doc.add_heading(f"{color} {heading}", level=2)
        codes = sorted(by_cat[cat], key=lambda c: -len(per_code[c]))
        for code in codes:
            entries = per_code[code]
            distinct = len({e["invoice_id"] for e in entries})
            desc = ERROR_DESCRIPTIONS.get(code, {})

            # Code heading
            h = doc.add_heading(code, level=3)
            h_run = h.runs[0]
            h_run.font.size = Pt(11)
            h_run.font.color.rgb = RGBColor.from_string("1F4E79")

            # Metadata table
            meta_tbl = doc.add_table(rows=0, cols=2)
            meta_tbl.style = "Light List Accent 1"
            for label, value in [
                ("Finding count",      f"{len(entries):,}"),
                ("Affected invoices",  f"{distinct:,}"),
                ("Category",           desc.get("category", "")),
                ("Severity",           desc.get("severity", "")),
                ("What is happening",  desc.get("what", "")),
                ("Action required",    desc.get("action", "")),
            ]:
                row = meta_tbl.add_row()
                _set_cell_text(row.cells[0], label, bold=True)
                _set_cell_text(row.cells[1], value)

            # Examples table
            doc.add_paragraph().add_run("Examples:").bold = True
            ex_tbl = doc.add_table(rows=1, cols=5)
            ex_tbl.style = "Light Grid Accent 1"
            _add_shaded_row(ex_tbl, ["Unit", "Supplier", "Bill No", "PO", "Amount"])
            seen = set()
            sample_count = 0
            detail_lines: List[str] = []
            for e in entries:
                if e["invoice_id"] in seen:
                    continue
                seen.add(e["invoice_id"])
                info = enriched.get(e["invoice_id"], {})
                row = ex_tbl.add_row()
                _set_cell_text(row.cells[0], str(info.get("unit") or "?"), size=8)
                _set_cell_text(row.cells[1], str(info.get("supplier_name") or "?")[:40], size=8)
                _set_cell_text(row.cells[2], str(info.get("invoice_number") or "?"), size=8)
                _set_cell_text(row.cells[3], str(info.get("po_number") or "(none)"), size=8)
                _set_cell_text(row.cells[4], _format_amt(info.get("total_amount")), size=8)
                line_sfx = f" [line {e['line_seq']}]" if e.get("line_seq") else ""
                detail_lines.append(
                    f"\u2022 {info.get('invoice_number', '?')}{line_sfx}: {e['message']}"
                )
                sample_count += 1
                if sample_count >= 5:
                    break

            # Detail lines paragraph
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            for line in detail_lines:
                r = p.add_run(line + "\n")
                r.font.size = Pt(8)
                r.font.color.rgb = RGBColor.from_string("555555")
            doc.add_paragraph()  # spacer

    # -- 4. Summary email template -------------------------------------------
    doc.add_page_break()
    doc.add_heading("4. Summary Email Template", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "The block below can be copy-pasted into an email to the srimukha "
        "data team. Attach this document as backing detail."
    ).italic = True

    template = (
        "Subject: Bill Register Mar-26 — data quality findings / action required\n\n"
        "Hi team,\n\n"
        "We ran the latest Bill Register through our automated invoice-validation\n"
        f"pipeline. Out of {total:,} invoices:\n\n"
        f"    • {buckets.get('validated', 0):>5} ({(buckets.get('validated', 0) / total * 100):.1f}%) validated cleanly — ready for payment flow.\n"
        f"    • {buckets.get('waiting_for_re_validation', 0):>5} ({(buckets.get('waiting_for_re_validation', 0) / total * 100):.1f}%) are legitimate shortfalls — debit-note flow.\n"
        f"    • {buckets.get('exception_approval', 0):>5} ({(buckets.get('exception_approval', 0) / total * 100):.1f}%) hit 'PO already fulfilled' — exception approval.\n"
        f"    • {buckets.get('waiting_for_validation', 0):>5} ({(buckets.get('waiting_for_validation', 0) / total * 100):.1f}%) cannot yet be validated due to source data issues.\n\n"
        "Four priority fixes will unblock most of the stuck invoices:\n\n"
        f"  1. Include subcontract orders (SC prefix) in the PO export.\n"
        f"     {_count('E003_PO_NOT_FOUND')} invoices cannot be validated because their PO is missing.\n\n"
        f"  2. Populate 'PO / SCO No.' on every invoice row in the Bill Register.\n"
        f"     {_count('E002_NO_PO_LINK')} invoices have no PO reference at all.\n\n"
        f"  3. Correct GST type on ~{_count('E034_INTRA_STATE_WITH_IGST') + _count('E035_INTER_STATE_WITH_CGST_SGST')} invoice lines where intra/inter-state classification\n"
        "     was wrong (compliance risk).\n\n"
        f"  4. Review price drift on {_count('E022_LINE_RATE_MISMATCH') + _count('E023_LINE_PRICE_MISMATCH')} invoice lines — rate does not match PO.\n\n"
        "Full per-category details, affected invoices, and real examples are in\n"
        "the attached document.\n\n"
        "Thanks."
    )
    p = doc.add_paragraph()
    run = p.add_run(template)
    run.font.size = Pt(10)
    run.font.name = "Consolas"

    # -- 5. Footer / about ----------------------------------------------------
    doc.add_page_break()
    doc.add_heading("5. About this report", level=1)
    about = doc.add_paragraph()
    about.add_run(
        "This report was generated by the automated email_automation validation "
        "engine. The engine re-runs every day after the Bill Register arrives, "
        "so the findings reflect the latest state of the srimukha data. The same "
        "engine evaluates portal-uploaded invoices and email-ingested invoices "
        "identically.\n\n"
    )
    about.add_run(
        "Error codes have the form ENNN_NAME (errors) or WNNN_NAME (warnings). "
        "Stable codes let the same finding be tracked across runs and make it "
        "easy to target bulk fixes at the source."
    )

    doc.save(str(DOCX_PATH))
    return DOCX_PATH


def try_pdf(docx_path: Path) -> Path:
    """Best-effort PDF conversion. Returns the path if successful, else None."""
    try:
        import docx2pdf
    except ImportError:
        print("  docx2pdf not installed — skipping PDF (install with: pip install docx2pdf)")
        return None
    try:
        docx2pdf.convert(str(docx_path), str(PDF_PATH))
        return PDF_PATH
    except Exception as exc:
        print(f"  PDF conversion failed: {exc}")
        return None


def save_json(data: Dict[str, Any]) -> None:
    JSON_PATH.write_text(json.dumps(data, indent=2, default=str), encoding="utf-8")


def main() -> int:
    setup_logging()
    print(f"Phase 3 doc generator — writing to {REPORT_DIR}")
    print()
    data = collect_data()

    print(f"\nSaving raw data to {JSON_PATH.name} ...")
    save_json(data)

    print(f"Building docx -> {DOCX_PATH.name} ...")
    build_docx(data)
    print(f"  docx ready ({DOCX_PATH.stat().st_size // 1024} KB)")

    print(f"Attempting PDF conversion ...")
    pdf = try_pdf(DOCX_PATH)
    if pdf:
        print(f"  pdf ready: {pdf} ({pdf.stat().st_size // 1024} KB)")
    else:
        print(f"  pdf not generated — docx is the primary deliverable")

    print()
    print("Files produced:")
    print(f"  {DOCX_PATH}")
    if pdf:
        print(f"  {pdf}")
    print(f"  {JSON_PATH}")
    close_pool()
    return 0


if __name__ == "__main__":
    sys.exit(main())
